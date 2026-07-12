from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from knowledge_assistant.domain.documents import NormalizedDocument
from knowledge_assistant.infrastructure.parsers.base import (
    DocumentParseError,
    DocumentParser,
    normalized_document,
)


class DocxParser(DocumentParser):
    supported_suffixes = (".docx",)

    def parse(self, path: Path) -> list[NormalizedDocument]:
        try:
            document = Document(path)
        except (BadZipFile, OSError, PackageNotFoundError, ValueError) as error:
            raise DocumentParseError(path, "could not parse DOCX file") from error

        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        title = next(
            (
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip() and paragraph.style.name.startswith("Heading")
            ),
            path.stem,
        )
        return [
            normalized_document(
                path=path,
                source_type="docx",
                content="\n\n".join(paragraphs),
                title=title,
            )
        ]
